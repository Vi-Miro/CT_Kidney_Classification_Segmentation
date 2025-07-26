import sys
import json
from PyQt5.QtWidgets import (
    QApplication, QWidget, QLabel, QPushButton, QFileDialog, QTextEdit,
    QVBoxLayout, QMessageBox, QComboBox, QGroupBox, QHBoxLayout
)
from PyQt5.QtGui import QPixmap, QImage
from PyQt5.QtCore import Qt
import pydicom
import numpy as np
import cv2
from classify_image import classify_dicom, set_model, get_current_model_info
from confluent_kafka import Producer, KafkaException, Consumer
import json
import os
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'  # Отключает большинство предупреждений TensorFlow
import absl.logging
absl.logging.set_verbosity(absl.logging.ERROR)
from kidney_detector import KidneyDetector
import base64
from datetime import datetime


class KafkaManager:
    def __init__(self, bootstrap_servers='localhost:9092'):
        self.bootstrap_servers = bootstrap_servers
        self.producer_config = {
            'bootstrap.servers': bootstrap_servers,
            'message.max.bytes': 20971520,
            'queue.buffering.max.messages': 100000,
            'queue.buffering.max.kbytes': 102400,
            'batch.num.messages': 1000,
            'message.timeout.ms': 5000,
            'enable.idempotence': True
        }

    def check_connection(self):
        """Проверка подключения к Kafka с использованием confluent-kafka"""
        try:
            consumer = Consumer({
                'bootstrap.servers': self.bootstrap_servers,
                'group.id': 'test-connection-group',
                'auto.offset.reset': 'earliest'
            })
            topics = consumer.list_topics().topics
            print(f"Available topics: {list(topics.keys())}")
            consumer.close()
            return True
        except Exception as e:
            print(f"Kafka connection error: {str(e)}")
            return False

    def delivery_report(self, err, msg):
        """Callback для отслеживания доставки с переносами строк"""
        if err is not None:
            print(f'\nMessage delivery failed: {err}\n')
        else:
            print(f'\nMessage delivered to:')
            print(f'  Topic:    {msg.topic()}')
            print(f'  Partition: {msg.partition()}')
            print(f'  Offset:    {msg.offset()}\n')

    def send_message(self, data, topic="TestTopic"):
        """Отправка сообщения в Kafka"""
        producer = Producer(self.producer_config)
        message = json.dumps(data, ensure_ascii=False)

        try:
            producer.produce(
                topic,
                value=message.encode('utf-8'),
                callback=self.delivery_report
            )
            producer.flush(timeout=5)
            return True, "Message sent successfully"
        except KafkaException as e:
            return False, f"Kafka error: {str(e)}"
        except Exception as e:
            return False, f"General error: {str(e)}"
            
    def send_dicom_file(self, dicom_path, result_data, topic="DICOMFiles"):
        """Отправка DICOM файла и результатов анализа в Kafka"""
        try:
            # Читаем DICOM файл как бинарные данные
            with open(dicom_path, 'rb') as f:
                dicom_bytes = f.read()
            
            # Кодируем в base64 для передачи через JSON
            dicom_base64 = base64.b64encode(dicom_bytes).decode('utf-8')
            
            # Создаем сообщение с метаданными и файлом
            message = {
                'metadata': result_data,
                'dicom_file': dicom_base64,
                'filename': os.path.basename(dicom_path)
            }
            
            producer = Producer(self.producer_config)

            print("Формируется сообщение...")
            producer.produce(
                topic,
                value=json.dumps(message).encode('utf-8'),
                callback=self.delivery_report
            )
            print("Отправка сообщения в Kafka...")
            producer.flush(timeout=10)
            print("Завершена отправка.")
            return True, "DICOM file sent successfully"
            
        except Exception as e:
            return False, f"Error sending DICOM file: {str(e)}"

def get_dicom_tags(dicom_file):
    """Извлекает основные теги из DICOM файла и возвращает их в виде словаря"""
    tags = {}
    try:
        ds = pydicom.dcmread(dicom_file)
        ds.decode()

        def safe_str(value):
            try:
                return str(value)
            except:
                return "Не удалось декодировать значение"
        
        # Основные теги пациента
        tags['PatientName'] = str(getattr(ds, 'PatientName', ''))
        tags['PatientID'] = str(getattr(ds, 'PatientID', ''))
        tags['PatientBirthDate'] = str(getattr(ds, 'PatientBirthDate', ''))
        tags['PatientSex'] = str(getattr(ds, 'PatientSex', ''))
        
        # Информация об исследовании
        tags['StudyDate'] = str(getattr(ds, 'StudyDate', ''))
        tags['StudyDescription'] = str(getattr(ds, 'StudyDescription', ''))
        tags['Modality'] = str(getattr(ds, 'Modality', ''))
        tags['BodyPartExamined'] = str(getattr(ds, 'BodyPartExamined', ''))
        
        # Теги изображения
        tags['Rows'] = str(getattr(ds, 'Rows', ''))
        tags['Columns'] = str(getattr(ds, 'Columns', ''))
        tags['PixelSpacing'] = str(getattr(ds, 'PixelSpacing', ''))
        tags['BitsAllocated'] = str(getattr(ds, 'BitsAllocated', ''))

        if hasattr(ds, 'Prediction'):
            tags['prediction'] = str(ds.Prediction)
            tags['confidence'] = str(ds.Confidence)
            tags['probabilities'] = str(ds.Probabilities)
        else:
            tags['prediction'] = "анализ не проводился"
            tags['confidence'] = "анализ не проводился"
            tags['probabilities'] = "анализ не проводился"
        
        
        # Конвертируем все значения в строки (для JSON сериализации)
        tags = {k: str(v) for k, v in tags.items()}
        
    except Exception as e:
        print(f"Ошибка при чтении DICOM тегов: {str(e)}")
        tags['error'] = str(e)
    
    return tags


def add_classification_result_to_dicom(dicom_path, result, output_path=None):
    try:
        ds = pydicom.dcmread(dicom_path)
        
        # Добавляем результаты как стандартные приватные теги
        ds.add_new(0x00231010, 'LO', result.get('prediction', ''))  # Prediction
        ds.add_new(0x00231011, 'DS', str(result.get('confidence', 0.0)))  # Confidence как Decimal String
        ds.add_new(0x00231012, 'LT', str(result.get('probabilities', {})))  # Long Text
        
        # Добавляем информацию о модели
        model_info = result.get('model_info', {})
        ds.add_new(0x00231020, 'LO', model_info.get('model_name', ''))
        ds.add_new(0x00231021, 'LO', model_info.get('model_path', ''))
        
        # Сохраняем файл
        save_path = output_path if output_path else dicom_path
        ds.save_as(save_path)
        return save_path
        
    except Exception as e:
        print(f"Ошибка при добавлении результатов в DICOM: {str(e)}")
        return None


class KidneyApp(QWidget):
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Классификация патологий почек')
        self.resize(850, 800)

        self.setStyleSheet("""
            QWidget {
                background-color: #2f2f2f;
                color: #f0f0f0;
                font-family: Segoe UI;
                font-size: 14px;
            }
            QLabel#ImageLabel {
                border: 2px dashed #555;
                background-color: #3b3b3b;
                color: #888;
                min-height: 300px;
                padding: 10px;
                qproperty-alignment: AlignCenter;
                font-size: 16px;
            }
            QPushButton {
                background-color: #4a90e2;
                color: white;
                border-radius: 8px;
                padding: 8px;
            }
            QPushButton:hover {
                background-color: #357ab8;
            }
            QTextEdit {
                background-color: #1e1e1e;
                border: 1px solid #555;
                border-radius: 6px;
                padding: 6px;
                color: #dcdcdc;
            }
        """)

        # Layout
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # DICOM
        self.label_image = QLabel("Изображение не загружено")
        self.label_image.setObjectName("ImageLabel")

        self.button_load = QPushButton("📂 Загрузить DICOM")
        self.button_tags = QPushButton("🔍 Показать DICOM теги")
        self.button_classify = QPushButton("🧠 Запустить анализ")
        self.button_save_dicom = QPushButton("💾 Сохранить DICOM с результатами")
        self.result_field = QTextEdit()
        self.save_button = QPushButton("💾 Сохранить результат в JSON")

        #self.button_save_image = QPushButton("💾 Сохранить изображение с результатами")
        #self.button_save_image.clicked.connect(self.save_image_result)
        #self.layout.addWidget(self.button_save_image)

        self.model_selector = QComboBox()
        self.model_selector.addItems(['VGG16_v1', 'CNN_v1'])
        self.model_selector.setCurrentText('CNN_v1')
        self.model_selector.currentTextChanged.connect(self.change_model)

        # Pack
        self.layout.addWidget(self.label_image)
        self.layout.addWidget(self.button_load)
        self.layout.addWidget(self.button_tags)
        self.layout.addWidget(self.button_classify)
        self.layout.addWidget(self.button_save_dicom)
        self.layout.addWidget(self.result_field)
        self.layout.addWidget(self.save_button)

        # Events
        self.button_load.clicked.connect(self.load_dicom)
        self.layout.insertWidget(2, QLabel("Выберите модель:"))
        self.layout.insertWidget(3, self.model_selector)
        self.button_tags.clicked.connect(self.show_dicom_tags)
        self.button_classify.clicked.connect(self.run_classification)
        self.button_save_dicom.clicked.connect(self.save_dicom_with_results)
        self.save_button.clicked.connect(self.save_result)

        self.button_save_report = QPushButton("📄 Сохранить медицинский отчет")
        self.button_save_report.clicked.connect(self.save_medical_report)
        self.layout.addWidget(self.button_save_report)

        self.kafka_button = QPushButton("Отправить результаты анализа в kafka")
        self.kafka_button.setStyleSheet("background-color: #2e8b57; color: white; font-weight: bold; padding: 6px;")
        self.kafka_button.clicked.connect(self.send_kafka_message)
        self.layout.addWidget(self.kafka_button)

        self.dicom_array = None
        self.result = None
        self.current_dicom_path = None
        self.kafka_manager = KafkaManager()

        self.show_disclaimer()

    def show_disclaimer(self):
        msg = QMessageBox(self)
        msg.setWindowTitle("Дисклеймер")
        msg.setIcon(QMessageBox.Warning)
        msg.setText("⚠️ Внимание!\n\nЭта программа предназначена только для образовательных целей и не может быть использована для медицинской диагностики.")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

    def generate_medical_report(self, result):
        """Генерация профессионального медицинского отчета"""
        prediction = result.get("prediction", "Unknown")
        confidence = result.get("confidence", 0)
        probabilities = result.get("probabilities", {})
        
        # Форматирование даты и времени
        current_time = datetime.now().strftime("%d.%m.%Y %H:%M")
        
        # Базовый шаблон отчета
        report = f"""
        МЕДИЦИНСКИЙ ОТЧЕТ О РЕЗУЛЬТАТАХ АНАЛИЗА
        Дата и время исследования: {current_time}
        Исследуемая область: Почки
        
        РЕЗУЛЬТАТЫ:
        """
        # Детализация результатов
        if prediction == "Normal":
            report += """
        - Патологических изменений не выявлено.
        """
        else:
            report += f"""
        - Обнаружены патологические изменения: {self.get_medical_term(prediction)}
        - Уверенность алгоритма: {confidence*100:.1f}%
        """
        
        # Добавляем полную вероятностную статистику
        report += "\nВЕРОЯТНОСТНАЯ ОЦЕНКА:\n"
        for class_name, prob in probabilities.items():
            report += f"- {self.get_medical_term(class_name)}: {prob*100:.1f}%\n"
        
        # Заключение и рекомендации
        report += """
        
        ЗАКЛЮЧЕНИЕ:
        """
        
        if prediction == "Normal":
            report += "КТ признаков патологии почек не выявлено."
        else:
            report += f"Выявленные КТ признаки: {self.get_medical_term(prediction)}."
        
        report += """
        
        РЕКОМЕНДАЦИИ:
        - Результаты автоматического анализа требуют интерпретации врачом-специалистом
        - При наличии клинических симптомов рекомендована консультация нефролога
        - Возможно потребуется дополнительное обследование
        
        Врач-радиолог: _________________________
        """
        
        return report.strip()

    def get_medical_term(self, class_name):
        """Преобразует названия классов в медицинские термины"""
        terms = {
            "Normal": "норма",
            "Cyst": "киста",
            "Stone": "камень",
            "Tumor": "опухоль"
        }
        return terms.get(class_name, class_name)

    def save_medical_report(self):
        """Сохранение отчета в файл"""
        if not hasattr(self, 'result') or not self.result:
            QMessageBox.warning(self, "Ошибка", "Сначала выполните анализ изображения.")
            return
        
        # Генерация отчета
        report = self.generate_medical_report(self.result)
        
        # Диалог сохранения файла
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить медицинский отчет",
            "",
            "Текстовые файлы (*.txt);;Документы Word (*.docx);;PDF файлы (*.pdf)"
        )
        
        if not path:
            return
        
        try:
            # Сохранение в выбранном формате
            if path.endswith('.txt'):
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(report)
            elif path.endswith('.docx'):
                from docx import Document
                doc = Document()
                doc.add_heading('Медицинский отчет', level=1)
                for line in report.split('\n'):
                    doc.add_paragraph(line.strip())
                doc.save(path)
            elif path.endswith('.pdf'):
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
                pdf.set_font('DejaVu', '', 12)
                for line in report.split('\n'):
                    pdf.cell(0, 10, line.strip(), ln=1)
                pdf.output(path)
            else:
                # По умолчанию сохраняем как txt
                with open(path + '.txt', 'w', encoding='utf-8') as f:
                    f.write(report)
            
            QMessageBox.information(self, "Успех", "Отчет успешно сохранен.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить отчет: {str(e)}")

    def run_classification(self):
        if self.dicom_array is not None:
            try:
                # Выполняем классификацию
                self.result = classify_dicom(self.dicom_array)
                
                if self.result["confidence"] == 0.0:
                    self.result_field.setText("⚠ Ошибка: модель не загружена.")
                else:
                    # 1. Записываем ТОЛЬКО технические данные в DICOM
                    if self.current_dicom_path and self.current_dicom_path.lower().endswith('.dcm'):
                        self.save_technical_results_to_dicom()
                    
                    # 2. Показываем полный медицинский отчет
                    report = self.generate_medical_report(self.result)
                    self.result_field.setText(report)
                    
                    # 3. Дополнительно показываем отчет в диалоговом окне
                    QMessageBox.information(
                        self, 
                        "Результаты анализа", 
                        report,
                        QMessageBox.Ok
                    )
                    
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка классификации: {str(e)}")



    
    def change_model(self, model_name):
        success = set_model(model_name)
        if success:
            QMessageBox.information(self, "Модель изменена", f"Текущая модель: {model_name}")
        else:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить модель {model_name}")

    def load_dicom(self):
        path, _ = QFileDialog.getOpenFileName(
            self, 
            "Открыть медицинское изображение", 
            "", 
            "DICOM Files (*.dcm);;JPEG Images (*.jpg *.jpeg);;PNG Images (*.png)"
        )
        if not path:
            return
        
        self.current_dicom_path = path
        
        try:
            # Загрузка изображения
            if path.lower().endswith('.dcm'):
                ds = pydicom.dcmread(path)
                self.dicom_array = ds.pixel_array
                if self.dicom_array.dtype != np.uint8:
                    self.dicom_array = cv2.convertScaleAbs(
                        self.dicom_array, 
                        alpha=255.0/self.dicom_array.max()
                    )
            else:
                self.dicom_array = cv2.imread(path, cv2.IMREAD_UNCHANGED)
            
            # Конвертация в RGB
            if len(self.dicom_array.shape) == 2:
                display_image = cv2.cvtColor(self.dicom_array, cv2.COLOR_GRAY2RGB)
            else:
                display_image = self.dicom_array.copy()
                if display_image.shape[2] == 4:
                    display_image = cv2.cvtColor(display_image, cv2.COLOR_RGBA2RGB)
                else:
                    display_image = cv2.cvtColor(display_image, cv2.COLOR_BGR2RGB)
            
            # Детекция почек
            detector = KidneyDetector('C:/Phyton/Проект/best.pt')
            detections = detector.detect(display_image)
            
            # Проверка наличия почек (теперь ищем left_kidney и right_kidney)
            kidney_found = any(d['class'] in ['left_kidney', 'right_kidney', 'kidney'] for d in detections)
            
            if not kidney_found:
                reply = QMessageBox.question(
                    self,
                    "Почки не обнаружены",
                    "На изображении не обнаружены почки. Продолжить анализ?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                
                if reply == QMessageBox.No:
                    self.reset_interface()
                    return
            
            # Отображение bounding boxes
            for det in detections:
                if det['class'] in ['left_kidney', 'right_kidney', 'kidney']:  # Добавили все возможные варианты
                    x1, y1, x2, y2 = det['bbox']
                    color = (0, 255, 0)  # Зеленый
                    if det['class'] == 'left_kidney':
                        color = (0, 255, 255)  # Желтый для левой почки
                    elif det['class'] == 'right_kidney':
                        color = (255, 255, 0)  # Голубой для правой почки
                    
                    cv2.rectangle(display_image, (x1, y1), (x2, y2), color, 2)
                    cv2.putText(display_image, 
                            f"{det['class']} {det['confidence']:.2f}", 
                            (x1, y1-10), 
                            cv2.FONT_HERSHEY_SIMPLEX, 
                            0.5, color, 1)
            
            # Отображение изображения
            self.display_image(display_image)
            
        except Exception as e:
            self.show_error(f"Не удалось загрузить изображение: {str(e)}")
            self.reset_interface()

    def reset_interface(self):
        self.dicom_array = None
        self.label_image.setText("Изображение не загружено")
        self.label_image.setPixmap(QPixmap())

    def display_image(self, image):
        h, w, ch = image.shape
        bytes_per_line = ch * w
        qimage = QImage(image.data, w, h, bytes_per_line, QImage.Format_RGB888)
        pixmap = QPixmap.fromImage(qimage).scaled(512, 512, Qt.KeepAspectRatio)
        self.label_image.setPixmap(pixmap)
        self.label_image.setText("")

    def show_error(self, message):
        print(message)
        QMessageBox.critical(self, "Ошибка", message)

    def show_dicom_tags(self):
        if not self.current_dicom_path:
            QMessageBox.warning(self, "Ошибка", "Сначала загрузите DICOM файл.")
            return
        
        # Сначала показываем результаты из памяти, если они есть
        if hasattr(self, 'result') and self.result:
            tags = {
                **get_dicom_tags(self.current_dicom_path),
                "prediction": str(self.result.get("prediction", "")),
                "confidence": str(self.result.get("confidence", "")),
                "probabilities": str(self.result.get("probabilities", ""))
            }
        else:
            # Если результатов в памяти нет, читаем из файла
            tags = get_dicom_tags(self.current_dicom_path)
        
        self.result_field.setText(json.dumps(tags, indent=4, ensure_ascii=False))

    def run_classification(self):
        if self.dicom_array is not None:
            try:
                # Выполняем классификацию
                self.result = classify_dicom(self.dicom_array)
                
                if self.result["confidence"] == 0.0:
                    self.result_field.setText("⚠ Ошибка: модель не загружена.")
                else:
                    # Сохраняем результаты в файл (если это DICOM)
                    if self.current_dicom_path and self.current_dicom_path.lower().endswith('.dcm'):
                        add_classification_result_to_dicom(
                            self.current_dicom_path, 
                            self.result,
                            self.current_dicom_path  # сохраняем в тот же файл
                        )
                    
                    # Показываем результаты
                    self.result_field.setText(json.dumps(self.result, indent=4, ensure_ascii=False))
                    
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка классификации: {str(e)}")
                print(f"Ошибка классификации: {str(e)}")

    def save_dicom_with_results(self):
        if self.current_dicom_path and self.result:
            output_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Сохранить DICOM с результатами", 
                "", 
                "DICOM Files (*.dcm)"
            ) 
            
            if output_path:
                saved_path = add_classification_result_to_dicom(
                    self.current_dicom_path, 
                    self.result, 
                    output_path
                )
                
                if saved_path:
                    QMessageBox.information(
                        self, 
                        "Сохранение", 
                        f"DICOM файл с результатами сохранен:\n{saved_path}"
                    )
                else:
                    QMessageBox.critical(
                        self, 
                        "Ошибка", 
                        "Не удалось сохранить DICOM файл с результатами"
                    )
        else:
            QMessageBox.warning(
                self, 
                "Ошибка", 
                "Сначала загрузите DICOM файл и выполните классификацию"
            )

    def send_kafka_message(self):
        if not self.current_dicom_path:
            QMessageBox.warning(self, "Ошибка", "Сначала загрузите DICOM файл")
            return
            
        if not self.result:
            QMessageBox.warning(self, "Ошибка", "Сначала выполните анализ изображения")
            return
            
        # Проверяем соединение с Kafka
        if not self.kafka_manager.check_connection():
            QMessageBox.critical(self, "Ошибка", "Не удалось подключиться к Kafka серверу")
            return
            
        # Получаем метаданные без "анализ не проводился"
        metadata = get_dicom_tags(self.current_dicom_path)
        metadata.update({
            'analysis_result': self.result,
            'model_info': get_current_model_info()
        })
        
        # Удаляем ненужные поля
        for field in ['prediction', 'confidence', 'probabilities']:
            metadata.pop(field, None)  # Удаляем, если они есть
            
        message = {'metadata': metadata}
        
        # Отправляем JSON с метаданными
        success, msg = self.kafka_manager.send_message(message, "AnalysisResults")
        if not success:
            QMessageBox.critical(self, "Ошибка", f"Не удалось отправить JSON: {msg}")
            return
            
        # Если это DICOM файл, отправляем и сам файл
        if self.current_dicom_path.lower().endswith('.dcm'):
            success, msg = self.kafka_manager.send_dicom_file(
                self.current_dicom_path,
                message['metadata'],
                "DICOMFiles"
            )
            
            if success:
                QMessageBox.information(self, "Успех", "DICOM файл и результаты анализа успешно отправлены")
            else:
                QMessageBox.warning(self, "Предупреждение", 
                    f"Результаты анализа отправлены, но не удалось отправить DICOM файл: {msg}")
        else:
            QMessageBox.information(self, "Успех", "Результаты анализа успешно отправлены")

    def save_result(self):
        if self.result:
            path, _ = QFileDialog.getSaveFileName(
                self, 
                "Сохранить JSON", 
                "classification_result.json", 
                "JSON (*.json)"
            )
            if path:
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(self.result, f, indent=4, ensure_ascii=False)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = KidneyApp()
    window.show()
    sys.exit(app.exec_())